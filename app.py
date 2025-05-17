import os
import sys
from flask import Flask, request, jsonify, render_template, send_from_directory, send_file, url_for, redirect, session
from flask_cors import CORS
from dotenv import load_dotenv
import google.genai as genai
from pymongo import MongoClient
from pymongo.server_api import ServerApi
from bson import ObjectId
import json
from authlib.integrations.flask_client import OAuth
import PyPDF2
import docx
from docx.shared import Pt
import datetime
import re
import io
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from urllib.parse import urlencode
import uuid
import random

# Import database logging functions
try:
    from init_database import add_system_log, log_user_activity
except ImportError:
    # Fallback logging functions if import fails
    def add_system_log(message, level="INFO"):
        print(f"[{level}] {message}")
    
    def log_user_activity(user_id, action, details=None):
        details_str = f" - {json.dumps(details)}" if details else ""
        print(f"[USER {user_id}] {action}{details_str}")

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__, static_folder='.')
CORS(app, supports_credentials=True)

# Set a more reliable secret key - Generate a fixed key for production
if os.getenv('SESSION_SECRET'):
    app.secret_key = os.getenv('SESSION_SECRET')
else:
    # Only for development - in production always use an environment variable
    app.secret_key = 'mr_wlah_dev_secret_key_12345'

# Configure session
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = datetime.timedelta(days=7)
app.config['SESSION_COOKIE_SECURE'] = False  # Set to True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_REFRESH_EACH_REQUEST'] = True
app.config['SESSION_USE_SIGNER'] = True

# Configure Google Gemini API
api_key = os.getenv('GEMINI_API_KEY')
genai_client = genai.Client(api_key=api_key)
model_name = os.getenv('GEMINI_MODEL', 'gemini-2.0-flash')

# Configure MongoDB - explicitly set the MongoDB URI for BenchAI
# This ensures we don't use any potentially incorrect values from .env
MONGODB_URI = "mongodb+srv://benchai.3cq4b8o.mongodb.net/?authSource=%24external&authMechanism=MONGODB-X509&retryWrites=true&w=majority&appName=MrWlah"
mongo_db = os.getenv('MONGODB_DATABASE', 'benchai')

# Initialize collections as None by default
users_collection = None
transformations_collection = None
api_usage_collection = None

# Check if X.509 certificate exists
cert_path = os.path.join('certs', 'X509-cert-5870665680541743449.pem')
cert_content = os.getenv('MONGODB_CERT')

if cert_content:
    # We're on Heroku with cert in env var
    import tempfile
    cert_file = tempfile.NamedTemporaryFile(delete=False)
    cert_file.write(cert_content.encode())
    cert_file.close()
    cert_path = cert_file.name
    has_certificate = True
else:
    # Local environment, check for file
    has_certificate = os.path.exists(cert_path)

# Connect to MongoDB if certificate exists
if has_certificate:
    try:
        print(f"Connecting to BenchAI MongoDB...")
        print(f"Using X.509 certificate at {cert_path}")
        
        # Set up MongoDB client with X.509 certificate
        mongo_client = MongoClient(
            MONGODB_URI,
            tls=True,
            tlsCertificateKeyFile=cert_path,
            server_api=ServerApi('1')
        )
        
        # Test connection
        mongo_client.admin.command('ping')
        print("✅ MongoDB connection successful")
        
        # Connect to the specified database
        db = mongo_client[mongo_db]
        add_system_log(f"Connected to MongoDB database: {mongo_db}")
        
        # Set up collections
        users_collection = db['users']
        transformations_collection = db['transformations']
        api_usage_collection = db['apiUsage']
        
        # Log connection success with database details
        collections = db.list_collection_names()
        add_system_log(f"Available collections: {', '.join(collections)}", "INFO")
        
    except Exception as e:
        error_msg = f"MongoDB connection error: {str(e)}"
        print(f"❌ {error_msg}")
        add_system_log(error_msg, "ERROR")
else:
    if not has_certificate:
        print(f"⚠️ X.509 certificate not found at {cert_path}")
        add_system_log(f"X.509 certificate not found at {cert_path}", "WARNING")
    
    print("Running in demo mode without database connection")
    add_system_log("Running in demo mode without database connection", "WARNING")

# Configure Auth0
oauth = OAuth(app)
auth0 = oauth.register(
    'auth0',
    client_id=os.getenv('AUTH0_CLIENT_ID'),
    client_secret=os.getenv('AUTH0_CLIENT_SECRET'),
    api_base_url=f"https://{os.getenv('AUTH0_DOMAIN')}",
    access_token_url=f"https://{os.getenv('AUTH0_DOMAIN')}/oauth/token",
    authorize_url=f"https://{os.getenv('AUTH0_DOMAIN')}/authorize",
    client_kwargs={
        'scope': 'openid profile email',
    },
    server_metadata_url=f"https://{os.getenv('AUTH0_DOMAIN')}/.well-known/openid-configuration"
)

# Helper function to clean LLM output
def clean_llm_response(text):
    """
    Removes common LLM prefacing and concluding meta-text from responses
    to provide only the usable transformed content.
    """
    # List of common prefacing patterns to remove
    prefacing_patterns = [
        r"^(Here'?s|Here is|I'?ve|I have|Below is|The following is).*?:\s*\n+",
        r"^(Sure|Okay|Alright|Of course|I'd be happy to|I can|I will).*?:\s*\n+",
        r"^(I'?ve transformed|I'?ve rewritten|I'?ve humanized|I'?ve modified).*?:\s*\n+",
        r"^(Your text|The text|This content) (has been|is now).*?:\s*\n+",
        r"^(In|With|Using|Employing|Applying) a.*?tone.*?:\s*\n+",
        r"^(As requested|As per your request|Based on your request).*?:\s*\n+",
        r"^(This is|Now the text is|Now it sounds) (more|much more|significantly).*?:\s*\n+",
        r"^(I've kept|While maintaining|Maintaining|I've maintained).*?:\s*\n+",
        r"^(Using|Incorporating|Adding|With) (personal|my own|human).*?:\s*\n+",
        r"^(Transformed version|Human version|Human-like version|Rewritten version).*?:\s*\n+",
    ]
    
    # List of common concluding patterns to remove
    concluding_patterns = [
        r"\n+\s*(I hope|Hope|Hopefully) (this|that|these|it).*?\.$",
        r"\n+\s*(Let me know|Feel free to|Please) (if|to|contact).*?\.$",
        r"\n+\s*(This|The text|This version|This rewrite) (should|now|has).*?\.$",
        r"\n+\s*(Is there|Do you|Would you|If you) (anything|like|need).*?\.$",
        r"\n+\s*(Thank you|Thanks) (for|and).*?\.$",
        r"\n+\s*(How'?s that|How does that sound|Does this work|What do you think).*?\.$",
        r"\n+\s*(I'?ve tried|I tried|I'?ve attempted) (to|my best).*?\.$",
        r"\n+\s*(I'?ve maintained|I maintained|I'?ve preserved) (the|your|original).*?\.$",
        r"\n+\s*(The word count|This keeps|I'?ve kept) (is|the|within).*?\.$",
        r"\n+\s*(This|The above|The text above) (maintains|keeps|preserves).*?\.$",
    ]
    
    # Apply all prefacing patterns
    for pattern in prefacing_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Apply all concluding patterns
    for pattern in concluding_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove any leading or trailing whitespace
    text = text.strip()
    
    return text

# Font style detection and preservation
def detect_font_style(text):
    """Detect font style markers in HTML or common text formatting"""
    font_info = {
        'font_family': None,
        'font_size': None,
        'font_style': None,
        'font_weight': None,
        'text_decoration': None,
        'color': None,
        'html_tags': []
    }
    
    # Check for HTML tags
    html_tags = re.findall(r'<([a-z0-9]+)[^>]*>', text, re.IGNORECASE)
    if html_tags:
        font_info['html_tags'] = list(set(html_tags))
    
    # Check for inline CSS
    font_family_match = re.search(r'font-family:\s*([^;]+);', text, re.IGNORECASE)
    if font_family_match:
        font_info['font_family'] = font_family_match.group(1).strip()
    
    font_size_match = re.search(r'font-size:\s*([^;]+);', text, re.IGNORECASE)
    if font_size_match:
        font_info['font_size'] = font_size_match.group(1).strip()
    
    # Check for direct font tags
    font_tag_match = re.search(r'<font[^>]+face=["\']([^"\']+)["\']', text, re.IGNORECASE)
    if font_tag_match:
        font_info['font_family'] = font_tag_match.group(1).strip()
    
    # Check for other style indicators
    if '<b>' in text.lower() or 'font-weight: bold' in text.lower():
        font_info['font_weight'] = 'bold'
    
    if '<i>' in text.lower() or 'font-style: italic' in text.lower():
        font_info['font_style'] = 'italic'
    
    if '<u>' in text.lower() or 'text-decoration: underline' in text.lower():
        font_info['text_decoration'] = 'underline'
    
    # Check for color
    color_match = re.search(r'color:\s*([^;]+);', text, re.IGNORECASE)
    if color_match:
        font_info['color'] = color_match.group(1).strip()
    
    return font_info

def apply_font_style(text, font_info):
    """Apply detected font style to the output text"""
    
    # If no style detected, return text as is
    if not any(font_info.values()):
        return text
    
    # If HTML tags detected, try to maintain structure
    if font_info['html_tags'] and not ('script' in font_info['html_tags'] or 'style' in font_info['html_tags']):
        # Apply basic styling
        styled_text = text
        
        # Apply font family if detected
        if font_info['font_family']:
            styled_text = f'<span style="font-family: {font_info["font_family"]}">{styled_text}</span>'
        
        # Apply font size if detected
        if font_info['font_size']:
            styled_text = f'<span style="font-size: {font_info["font_size"]}">{styled_text}</span>'
        
        # Apply bold if detected
        if font_info['font_weight'] == 'bold':
            styled_text = f'<b>{styled_text}</b>'
        
        # Apply italic if detected
        if font_info['font_style'] == 'italic':
            styled_text = f'<i>{styled_text}</i>'
        
        # Apply underline if detected
        if font_info['text_decoration'] == 'underline':
            styled_text = f'<u>{styled_text}</u>'
        
        # Apply color if detected
        if font_info['color']:
            styled_text = f'<span style="color: {font_info["color"]}">{styled_text}</span>'
        
        return styled_text
    
    return text

# Helper function to extract text from different file types
def extract_text_from_file(file):
    if file.filename.endswith('.txt'):
        return file.read().decode('utf-8')
    
    elif file.filename.endswith('.pdf'):
        try:
            # Enhanced PDF text extraction
            reader = PyPDF2.PdfReader(file)
            total_pages = len(reader.pages)
            text = []
            
            # Extract text from each page with structure preservation
            for i in range(total_pages):
                page = reader.pages[i]
                page_text = page.extract_text()
                
                # Basic structure preservation
                if page_text:
                    # Add page number for longer documents
                    if total_pages > 1:
                        text.append(f"--- Page {i+1} ---\n")
                    
                    # Clean up common PDF extraction issues
                    lines = page_text.split('\n')
                    cleaned_lines = []
                    
                    for line in lines:
                        # Remove excessive spaces
                        cleaned_line = re.sub(r'\s+', ' ', line).strip()
                        if cleaned_line:
                            cleaned_lines.append(cleaned_line)
                    
                    # Join with proper paragraph breaks
                    text.append('\n'.join(cleaned_lines))
                    
                    # Add extra newline between pages
                    if i < total_pages - 1:
                        text.append("\n")
            
            return "\n".join(text)
        except Exception as e:
            add_system_log(f"PDF extraction error: {str(e)}", "ERROR")
            # Fallback to basic extraction
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
            
    elif file.filename.endswith('.docx'):
        try:
            # Enhanced DOCX text extraction
            doc = docx.Document(file)
            full_text = []
            
            # Process paragraphs with structure preservation
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n\n".join(full_text)
        except Exception as e:
            add_system_log(f"DOCX extraction error: {str(e)}", "ERROR")
            # Fallback to basic extraction
            doc = docx.Document(file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

# Custom JSON encoder to handle MongoDB ObjectId
class JSONEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, ObjectId):
            return str(o)
        return json.JSONEncoder.default(self, o)

app.json_encoder = JSONEncoder

# Routes
@app.route('/')
def index():
    # Check if user is logged in
    is_logged_in = 'logged_in' in session and session['logged_in'] == True
    
    # Add detailed logging to troubleshoot the issue
    if is_logged_in:
        # User is authenticated
        user_info = session.get('profile', {})
        add_system_log(f"[INDEX ROUTE] User accessing homepage: {user_info.get('name', 'Unknown')} with session ID: {id(session)}", "INFO")
        add_system_log(f"[INDEX ROUTE] Session data: logged_in={session.get('logged_in')}, has_profile={('profile' in session)}", "INFO")
        
        # Make sure session is persisted
        session.modified = True
        
        # Serve the main application
        return send_file('index.html')
    else:
        # Look for session but not properly logged in
        if 'profile' in session:
            add_system_log(f"[INDEX ROUTE] Session exists but not properly logged in, clearing session. Session ID: {id(session)}", "WARNING")
            add_system_log(f"[INDEX ROUTE] Session data before clearing: {dict(session)}", "WARNING")
            session.clear()
        
        add_system_log("[INDEX ROUTE] Unauthenticated user attempting to access homepage, redirecting to login", "INFO")
        
        # If not authenticated, redirect to login
        return redirect('/login')

@app.route('/login')
def login_page():
    # If user is already logged in, redirect to homepage
    if 'logged_in' in session and session['logged_in'] == True:
        add_system_log(f"[LOGIN ROUTE] Authenticated user accessing login page, redirecting to homepage. Session ID: {id(session)}", "INFO")
        add_system_log(f"[LOGIN ROUTE] Session data: {dict(session)}", "INFO")
        
        # Ensure session data persists
        session.modified = True
        
        return redirect('/')
    
    # Check if there's an error parameter
    error = request.args.get('error')
    if error:
        add_system_log(f"[LOGIN ROUTE] Login page accessed with error: {error}", "WARNING")
    
    # Log for debugging the double login issue
    add_system_log(f"[LOGIN ROUTE] Serving login page to unauthenticated user. Session ID: {id(session)}", "INFO")
    add_system_log(f"[LOGIN ROUTE] Current session data: {dict(session)}", "INFO")
    
    # Otherwise serve the login page
    return send_file('login.html')

@app.route('/logout')
def logout_page():
    return send_file('logout.html')

@app.route('/<path:path>')
def serve_static(path):
    # Check if this is an admin path first
    if path.startswith('admin/'):
        # Extract the file name from the path
        admin_file = path.replace('admin/', '', 1)
        return send_from_directory('admin', admin_file)
    
    # If not an admin path, serve from the root directory
    return send_from_directory('.', path)

@app.route('/api/document/status', methods=['GET'])
def document_processing_status():
    """Endpoint to check document processing status for large files"""
    job_id = request.args.get('job_id')
    
    if not job_id:
        return jsonify({
            'status': 'error',
            'message': 'No job ID provided'
        }), 400
    
    # For simplicity, we'll just return a success status
    # In a production app, you'd check a queue or database for the actual status
    return jsonify({
        'status': 'completed',
        'job_id': job_id,
        'progress': 100
    })

# User subscription configurations
SUBSCRIPTION_LIMITS = {
    'FREE': 2,
    'BASIC': 10,
    'PRO': 25,
    'UNLIMITED': float('inf')
}

# Admin PIN for admin panel access (should be stored in env var in production)
ADMIN_PIN = os.getenv('ADMIN_PIN', '123456')  # Default for development only

@app.route('/admin')
def admin_panel():
    """Serve the admin panel"""
    return send_file('admin.html')

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    """Admin login with PIN"""
    data = request.json
    pin = data.get('pin')
    
    if not pin or pin != ADMIN_PIN:
        add_system_log("Failed admin login attempt with incorrect PIN", "WARNING")
        return jsonify({'success': False, 'message': 'Invalid PIN'}), 401
    
    # Set admin session
    session['is_admin'] = True
    
    # Store the admin's user profile if available
    admin_profile = session.get('profile', {})
    admin_user_id = None
    
    if admin_profile:
        # If an authenticated user is becoming an admin, log it
        admin_email = admin_profile.get('email', 'Unknown')
        admin_user_id = admin_profile.get('user_id', 'Unknown')
        add_system_log(f"User {admin_email} ({admin_user_id}) logged in as admin", "INFO")
        
        # If we have a MongoDB connection, flag this user as an admin in the database
        if users_collection is not None and admin_user_id:
            try:
                # First try to find user by user_id field (primary identifier)
                user = users_collection.find_one({'user_id': admin_user_id})
                
                # If not found, try auth0Id as fallback
                if not user:
                    user = users_collection.find_one({'auth0Id': admin_user_id})
                
                if user:
                    # Update the user record to mark as admin
                    users_collection.update_one(
                        {'_id': user['_id']},
                        {'$set': {
                            'is_admin': True,
                            'lastAdminLogin': datetime.datetime.utcnow(),
                            'user_id': admin_user_id  # Ensure user_id is set correctly
                        }}
                    )
                    add_system_log(f"Updated user record for admin: {admin_email}", "INFO")
                else:
                    add_system_log(f"Admin user not found in database: {admin_email}", "WARNING")
            except Exception as e:
                add_system_log(f"Error updating admin user record: {str(e)}", "ERROR")
    else:
        add_system_log(f"Admin logged in successfully (no user profile)", "INFO")
    
    return jsonify({
        'success': True, 
        'profile': {
            'email': admin_profile.get('email', ''),
            'name': admin_profile.get('name', ''),
            'user_id': admin_user_id
        } if admin_profile else {}
    })

@app.route('/api/admin/status')
def admin_status():
    """Check if user is logged in as admin"""
    is_admin = session.get('is_admin', False)
    return jsonify({'isAdmin': is_admin})

@app.route('/api/admin/users')
def admin_get_users():
    """Get all users for admin panel, with optional filters for recent activity and by Auth0 ID/user_id"""
    # Check if admin
    if not session.get('is_admin', False):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        if users_collection is not None:
            add_system_log("Fetching users from database for admin panel", "INFO")
            admin_profile = session.get('profile', {})
            admin_user_id = admin_profile.get('user_id', '')

            # --- NEW: Support query params ---
            recent = request.args.get('recent', 'false').lower() == 'true'
            auth0_id = request.args.get('auth0_id')
            user_id_param = request.args.get('user_id')
            query = {}
            now = datetime.datetime.utcnow()
            
            if recent:
                # Only users active in the last 24 hours
                since = now - datetime.timedelta(hours=24)
                query['lastActive'] = {'$gte': since}
            if auth0_id:
                query['$or'] = [
                    {'auth0Id': auth0_id},
                    {'user_id': auth0_id}
                ]
            if user_id_param:
                query['$or'] = [
                    {'user_id': user_id_param},
                    {'auth0Id': user_id_param}
                ]
            
            # If both recent and id filter, combine
            # (MongoDB $and for both filters)
            if (recent and (auth0_id or user_id_param)):
                or_filter = query.pop('$or')
                query = {'$and': [query, {'$or': or_filter}]}
            
            users = list(users_collection.find(query).sort([
                ('lastLogin', -1), 
                ('lastActive', -1),
                ('createdAt', -1)
            ]))
            add_system_log(f"Found {len(users)} users in database (query: {query})", "INFO")
            for user in users:
                if users.index(user) == 0:
                    add_system_log(f"Sample user data: {user}", "INFO")
                user['_id'] = str(user['_id'])
                if 'user_id' not in user:
                    if 'auth0Id' in user:
                        user['user_id'] = user['auth0Id']
                    elif 'auth0_id' in user:
                        user['user_id'] = user['auth0_id']
                    else:
                        user['user_id'] = user['_id']
                if 'lastLogin' in user and 'lastActive' not in user:
                    user['lastActive'] = user['lastLogin']
                elif 'last_login' in user and 'lastActive' not in user:
                    user['lastActive'] = user['last_login']
                elif 'lastActive' not in user:
                    user['lastActive'] = user.get('createdAt', None) or user.get('created_at', None)
                if 'subscription' not in user:
                    user['subscription'] = 'FREE'
                if 'usageCount' not in user:
                    user['usageCount'] = 0
                if 'name' not in user and 'email' in user:
                    user['name'] = user['email'].split('@')[0]
                elif 'name' not in user:
                    user['name'] = 'Unknown User'
                if 'email' not in user:
                    user['email'] = 'No email'
                if admin_user_id and user.get('user_id') == admin_user_id:
                    user['isCurrentAdmin'] = True
            add_system_log(f"Admin fetched user list ({len(users)} users)", "INFO")
            if not users:
                add_system_log("No users found in database, returning sample data", "WARNING")
                if admin_profile and admin_profile.get('user_id'):
                    users.append({
                        '_id': 'sample-admin',
                        'name': admin_profile.get('name', 'Admin User'),
                        'email': admin_profile.get('email', 'admin@example.com'),
                        'subscription': 'UNLIMITED',
                        'usageCount': 0,
                        'lastActive': datetime.datetime.now().isoformat(),
                        'user_id': admin_profile.get('user_id'),
                        'isCurrentAdmin': True
                    })
                else:
                    users.append({
                        '_id': 'sample-1',
                        'name': 'Sample User',
                        'email': 'user@example.com',
                        'subscription': 'FREE',
                        'usageCount': 1,
                        'lastActive': datetime.datetime.now().isoformat(),
                        'user_id': 'sample-user-id'
                    })
            return jsonify(users)
        else:
            add_system_log("No database connection, returning sample data", "WARNING")
            admin_profile = session.get('profile', {})
            sample_data = []
            if admin_profile and admin_profile.get('user_id'):
                sample_data.append({
                    '_id': 'admin-id',
                    'name': admin_profile.get('name', 'Admin User'),
                    'email': admin_profile.get('email', 'admin@example.com'),
                    'subscription': 'UNLIMITED',
                    'usageCount': 0,
                    'lastActive': datetime.datetime.now().isoformat(),
                    'user_id': admin_profile.get('user_id'),
                    'isCurrentAdmin': True
                })
            sample_data.append({
                '_id': 'sample-1',
                'name': 'Sample User',
                'email': 'user@example.com',
                'subscription': 'FREE',
                'usageCount': 1,
                'lastActive': datetime.datetime.now().isoformat(),
                'user_id': 'sample-user-id'
            })
            return jsonify(sample_data)
    except Exception as e:
        add_system_log(f"Error fetching users: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/user/<user_id>/grant', methods=['POST'])
def admin_grant_subscription(user_id):
    """Grant subscription to user"""
    # Check if admin
    if not session.get('is_admin', False):
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.json
    subscription = data.get('subscription', '').upper()
    
    if subscription not in SUBSCRIPTION_LIMITS:
        return jsonify({'error': 'Invalid subscription type'}), 400
    
    try:
        if users_collection is not None:
            # First try to find user by user_id field (primary identifier)
            user = users_collection.find_one({'user_id': user_id})
            
            # If not found, try other identifiers as fallback
            if not user and ObjectId.is_valid(user_id):
                user = users_collection.find_one({'_id': ObjectId(user_id)})
            if not user:
                user = users_collection.find_one({'auth0Id': user_id})
            
            if not user:
                return jsonify({'error': 'User not found'}), 404
                
            # Update user subscription
            update_result = users_collection.update_one(
                {'_id': user['_id']},
                {'$set': {
                    'subscription': subscription,
                    'subscriptionUpdatedAt': datetime.datetime.utcnow(),
                    'user_id': user.get('auth0Id', user_id)  # Ensure user_id is set
                }}
            )
            
            # Log the update
            if update_result.modified_count > 0:
                add_system_log(f"Admin granted {subscription} subscription to user {user_id} (name: {user.get('name', 'Unknown')})", "INFO")
            else:
                add_system_log(f"Admin attempted to grant {subscription} subscription to user {user_id}, but no changes were made", "WARNING")
            
            # Reset usage count if upgrading subscription
            users_collection.update_one(
                {'_id': user['_id']},
                {'$set': {'usageCount': 0}}
            )
            
            return jsonify({'success': True})
        else:
            # Mock response for no DB
            return jsonify({'success': True})
    except Exception as e:
        add_system_log(f"Error granting subscription: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/user/<user_id>/revoke', methods=['POST'])
def admin_revoke_subscription(user_id):
    """Revoke subscription from user"""
    # Check if admin
    if not session.get('is_admin', False):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        if users_collection is not None:
            # First try to find user by user_id field (primary identifier)
            user = users_collection.find_one({'user_id': user_id})
            
            # If not found, try other identifiers as fallback
            if not user and ObjectId.is_valid(user_id):
                user = users_collection.find_one({'_id': ObjectId(user_id)})
            if not user:
                user = users_collection.find_one({'auth0Id': user_id})
            
            if not user:
                return jsonify({'error': 'User not found'}), 404
            
            # Update user subscription to FREE
            update_result = users_collection.update_one(
                {'_id': user['_id']},
                {'$set': {
                    'subscription': 'FREE',
                    'subscriptionUpdatedAt': datetime.datetime.utcnow(),
                    'usageCount': 0,  # Reset usage count
                    'user_id': user.get('auth0Id', user_id)  # Ensure user_id is set
                }}
            )
            
            # Log the update
            if update_result.modified_count > 0:
                add_system_log(f"Admin revoked subscription from user {user_id} (name: {user.get('name', 'Unknown')})", "INFO")
            else:
                add_system_log(f"Admin attempted to revoke subscription from user {user_id}, but no changes were made", "WARNING")
            
            return jsonify({'success': True})
        else:
            # Mock response for no DB
            return jsonify({'success': True})
    except Exception as e:
        add_system_log(f"Error revoking subscription: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/user', methods=['POST'])
def admin_add_user():
    """Add or update a user manually"""
    # Check if admin
    if not session.get('is_admin', False):
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.json
    
    # Validate required fields
    required_fields = ['user_id', 'email', 'name', 'subscription']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Validate subscription type
    subscription = data.get('subscription', '').upper()
    if subscription not in SUBSCRIPTION_LIMITS:
        return jsonify({'error': f'Invalid subscription type: {subscription}'}), 400
    
    try:
        if users_collection is not None:
            # Check if user already exists
            existing_user = users_collection.find_one({'user_id': data['user_id']})
            
            timestamp_now = datetime.datetime.utcnow()
            
            if existing_user:
                # Update existing user
                update_data = {
                    'email': data['email'],
                    'name': data['name'],
                    'subscription': subscription,
                    'subscriptionUpdatedAt': timestamp_now,
                    'lastActive': timestamp_now
                }
                
                # Add optional fields if provided
                if 'usageCount' in data:
                    update_data['usageCount'] = int(data['usageCount'])
                
                update_result = users_collection.update_one(
                    {'_id': existing_user['_id']},
                    {'$set': update_data}
                )
                
                # Log the update
                if update_result.modified_count > 0:
                    add_system_log(f"Admin updated user: {data['user_id']} ({data['name']})", "INFO")
                
                return jsonify({
                    'success': True,
                    'message': 'User updated successfully',
                    'user_id': data['user_id']
                })
            else:
                # Create new user
                
                # Generate a username if not provided
                username = data.get('username')
                if not username:
                    email_base = data['email'].split('@')[0].replace('.', '').lower()
                    username = f"{email_base}_{random.randint(1000, 9999)}"
                
                new_user = {
                    'user_id': data['user_id'],
                    'auth0Id': data['user_id'],  # Keep backward compatibility
                    'email': data['email'],
                    'name': data['name'],
                    'username': username,        # Always include a username
                    'createdAt': timestamp_now,
                    'lastLogin': timestamp_now,
                    'lastActive': timestamp_now,
                    'subscription': subscription,
                    'usageCount': int(data.get('usageCount', 0)),
                    'subscriptionUpdatedAt': timestamp_now,
                    'preferences': {
                        'defaultTone': 'casual',
                        'saveHistory': True
                    },
                    'manually_added': True
                }
                
                insert_result = users_collection.insert_one(new_user)
                
                add_system_log(f"Admin manually added user: {data['user_id']} ({data['name']})", "INFO")
                
                return jsonify({
                    'success': True,
                    'message': 'User added successfully',
                    'user_id': data['user_id'],
                    '_id': str(insert_result.inserted_id)
                })
        else:
            # Mock response for no DB
            return jsonify({
                'success': True,
                'demo': True,
                'message': 'User would be added/updated (demo mode)',
                'user_id': data['user_id']
            })
    except Exception as e:
        add_system_log(f"Error adding/updating user: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/user/subscription')
def get_user_subscription():
    """Get current user's subscription info"""
    # Get user profile from session
    profile = session.get('profile')
    
    if not profile:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_id = profile.get('user_id')
    
    try:
        if users_collection is not None:
            # Get user from database - prioritize user_id field
            user = users_collection.find_one({'user_id': user_id})
            
            # Fallback to auth0Id if not found by user_id
            if not user:
                user = users_collection.find_one({'auth0Id': user_id})
            
            if user:
                # Return subscription info
                subscription_data = {
                    'subscription': user.get('subscription', 'FREE'),
                    'usageCount': user.get('usageCount', 0),
                    'usageLimit': SUBSCRIPTION_LIMITS.get(user.get('subscription', 'FREE'), SUBSCRIPTION_LIMITS['FREE']),
                    'lastResetDate': user.get('lastResetDate'),
                    'user_id': user_id
                }
                
                # Add extra data for debugging
                add_system_log(f"Retrieved subscription info for user {user_id}: {subscription_data['subscription']}, {subscription_data['usageCount']}/{subscription_data['usageLimit']}", "INFO")
                
                return jsonify(subscription_data)
            else:
                # User not found, create a new record
                add_system_log(f"User {user_id} not found in database, creating new record", "INFO")
                
                # Create new user with default FREE subscription
                timestamp_now = datetime.datetime.now()
                new_user = {
                    'user_id': user_id,
                    'auth0Id': user_id,  # Store both for backward compatibility
                    'email': profile.get('email', ''),
                    'name': profile.get('name', ''),
                    'createdAt': timestamp_now,
                    'lastLogin': timestamp_now,
                    'lastActive': timestamp_now,
                    'subscription': 'FREE',
                    'usageCount': 0,
                    'subscriptionUpdatedAt': timestamp_now,
                    'preferences': {
                        'defaultTone': 'casual',
                        'saveHistory': True
                    }
                }
                
                users_collection.insert_one(new_user)
                
                # Return default subscription info
                return jsonify({
                    'subscription': 'FREE',
                    'usageCount': 0,
                    'usageLimit': SUBSCRIPTION_LIMITS['FREE'],
                    'lastResetDate': None,
                    'user_id': user_id
                })
                
        else:
            # Mock response for no DB
            return jsonify({
                'subscription': 'FREE',
                'usageCount': 0,
                'usageLimit': SUBSCRIPTION_LIMITS['FREE'],
                'lastResetDate': None,
                'user_id': user_id
            })
    except Exception as e:
        add_system_log(f"Error getting subscription info: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/user/record-transformation', methods=['POST'])
def record_transformation():
    """Record a text transformation for the current user"""
    # Get user profile from session
    profile = session.get('profile')
    
    if not profile:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_id = profile.get('user_id')
    
    try:
        # Get the timestamp from the request
        data = request.json or {}
        timestamp = data.get('timestamp', datetime.datetime.utcnow().isoformat())
        
        if users_collection is not None:
            # Prioritize looking for user by user_id
            user = users_collection.find_one({'user_id': user_id})
            
            # Fallback to auth0Id if not found
            if not user:
                user = users_collection.find_one({'auth0Id': user_id})
            
            if user:
                # Update existing user
                users_collection.update_one(
                    {'_id': user['_id']},
                    {
                        '$inc': {'usageCount': 1},
                        '$set': {
                            'lastActive': timestamp,
                            # Ensure user_id is always set
                            'user_id': user_id
                        }
                    }
                )
            else:
                # If user not found, create a new record
                add_system_log(f"User {user_id} not found in database for recording transformation, creating new record", "INFO")
                
                # Create new user with default FREE subscription
                timestamp_now = datetime.datetime.now()
                new_user = {
                    'user_id': user_id,
                    'auth0Id': user_id,  # Store both for backward compatibility
                    'email': profile.get('email', ''),
                    'name': profile.get('name', ''),
                    'createdAt': timestamp_now,
                    'lastLogin': timestamp_now,
                    'lastActive': timestamp_now,
                    'subscription': 'FREE',
                    'usageCount': 1,  # Start with the current transformation
                    'subscriptionUpdatedAt': timestamp_now,
                    'preferences': {
                        'defaultTone': 'casual',
                        'saveHistory': True
                    }
                }
                
                users_collection.insert_one(new_user)
            
            # Record transformation in history if enabled
            if transformations_collection is not None:
                transformations_collection.insert_one({
                    'user_id': user_id,
                    'timestamp': timestamp,
                    'type': 'text'
                })
            
            add_system_log(f"Recorded transformation for user {user_id}", "INFO")
            return jsonify({'success': True})
        else:
            # Mock response for no DB
            return jsonify({'success': True})
    except Exception as e:
        add_system_log(f"Error recording transformation: {str(e)}", "ERROR")
        return jsonify({'error': str(e)}), 500

@app.route('/api/transform', methods=['POST'])
def transform_text():
    """Transform text using Gemini"""
    # Get user profile from session
    profile = session.get('profile')
    
    if not profile:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_id = profile.get('user_id')
    
    # Check if user has reached usage limit
    if users_collection is not None:
        # First look for user by user_id
        user = users_collection.find_one({'user_id': user_id})
        
        # If not found, fallback to auth0Id
        if not user:
            user = users_collection.find_one({'auth0Id': user_id})
        
        if user:
            subscription = user.get('subscription', 'FREE')
            usage_count = user.get('usageCount', 0)
            
            # Get limit based on subscription
            usage_limit = SUBSCRIPTION_LIMITS.get(subscription, SUBSCRIPTION_LIMITS['FREE'])
            
            # Check if user has exceeded limit
            if usage_count >= usage_limit and subscription != 'UNLIMITED':
                add_system_log(f"User {user_id} exceeded usage limit ({usage_count}/{usage_limit})", "INFO")
                return jsonify({
                    'error': 'Usage limit exceeded',
                    'message': 'You have reached your usage limit. Please upgrade your subscription to continue.'
                }), 402  # 402 Payment Required
    
    try:
        # Check if this is a file upload from FormData
        if request.files and 'file' in request.files:
            file = request.files['file']
            try:
                # Extract text from the file
                text = extract_text_from_file(file)
                
                # If we're just extracting text for display, return it
                if request.form.get('extract_only') == 'true':
                    return jsonify({'originalText': text})
                
                # Otherwise, set it for transformation below
                tone = request.form.get('tone', 'casual')
                preserve_font = request.form.get('preserveFont', 'true') == 'true'
                target_word_count = request.form.get('targetWordCount')
                if target_word_count:
                    target_word_count = int(target_word_count)
            except Exception as e:
                return jsonify({'error': f"Error processing file: {str(e)}"}), 400
        else:
            # Get request data from JSON
            data = request.get_json() if request.is_json else {}
            text = data.get('text', '')
            tone = data.get('tone', 'casual')
            preserve_font = data.get('preserveFont', True)
            target_word_count = data.get('targetWordCount')
            
            if not text:
                return jsonify({'error': 'No text or file provided'}), 400
        
        # Log transformation request
        if user_id:
            log_details = {
                "tone": tone,
                "text_length": len(text),
                "preserve_font": preserve_font,
                "target_word_count": target_word_count
            }
            log_user_activity(user_id, "TRANSFORM_TEXT", log_details)
        
        # If this is just a file upload without immediate transformation
        if request.files and not request.form.get('transform', False):
            # Return the extracted text without transformation
            return jsonify({
                'originalText': text,
                'message': 'File processed successfully'
            })
        
        # Detect font style if preservation is requested
        font_info = detect_font_style(text) if preserve_font else {}
        
        # Calculate original word count if target requested
        original_word_count = len(text.split()) if target_word_count else None
        
        # Create prompt for Gemini, including word count constraint if specified
        if target_word_count:
            prompt = f"""Rewrite the following text to sound more human. 
                     Add personal experiences, vary sentence structure, 
                     use colloquialisms, and make it engaging.
                     Use a {tone} tone.
                     
                     IMPORTANT: The output must be approximately {target_word_count} words (±100 words).
                     Current word count is approximately {original_word_count} words.
                     
                     IMPORTANT: Do not include any introductory phrases like "Here's your transformed text:" 
                     or concluding phrases like "I hope this helps!". Just provide the transformed content directly.
                     
                     Text to transform: {text}"""
        else:
            prompt = f"""Rewrite the following text to sound more human. 
                     Add personal experiences, vary sentence structure, 
                     use colloquialisms, and make it engaging.
                     Use a {tone} tone.
                     
                     IMPORTANT: Do not include any introductory phrases like "Here's your transformed text:" 
                     or concluding phrases like "I hope this helps!". Just provide the transformed content directly.
                     
                     Text to transform: {text}"""
        
        # Call Gemini API with new client pattern
        response = genai_client.models.generate_content(
            model=model_name,
            contents=prompt
        )
        
        transformed_text = response.text
        
        # Clean the LLM response to remove any prefacing or concluding meta-text
        transformed_text = clean_llm_response(transformed_text)
        
        # Apply original font style if preservation is requested
        if preserve_font:
            transformed_text = apply_font_style(transformed_text, font_info)
        
        # Log the transformation if MongoDB is configured and user is authenticated
        if transformations_collection is not None and user_id:
            try:
                # Create the transformation record
                transformation = {
                    'userId': user_id,  # Use user_id as the primary identifier
                    'user_id': user_id,  # Also store as user_id for consistency with other collections
                    'originalText': text,
                    'transformedText': transformed_text,
                    'tone': tone,
                    'fontStylePreserved': preserve_font,
                    'createdAt': datetime.datetime.now(),
                    'metadata': {
                        'characterCount': len(text),
                        'wordCount': original_word_count,
                        'targetWordCount': target_word_count,
                        'sourceType': 'file' if request.files else 'paste',
                        'modelUsed': model_name
                    }
                }
                
                # Insert the transformation record
                result = transformations_collection.insert_one(transformation)
                
                # Log successful storage
                if result.inserted_id:
                    add_system_log(f"Transformation record stored with ID: {result.inserted_id}", "INFO")
            except Exception as db_error:
                # Log the error but don't fail the request
                add_system_log(f"Failed to store transformation: {str(db_error)}", "ERROR")
        
        return jsonify({
            'transformedText': transformed_text, 
            'fontInfo': font_info,
            'originalText': text
        })
    
    except Exception as e:
        error_msg = f"Error transforming text: {str(e)}"
        print(error_msg)
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': 'Failed to transform text'}), 500

@app.route('/api/user/transformations', methods=['GET'])
def get_user_transformations():
    if transformations_collection is None:
        return jsonify({'error': 'MongoDB not configured', 'demo': True}), 200
    
    user_id = request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'User ID required'}), 400
    
    try:
        # Find transformations for this user by searching for both userId and user_id fields
        transformations = list(transformations_collection.find(
            {'$or': [
                {'userId': user_id},  # Backward compatibility with older records
                {'user_id': user_id}   # New user_id field
            ]}
        ).sort('createdAt', -1).limit(10))
        
        add_system_log(f"Retrieved {len(transformations)} transformations for user {user_id}", "INFO")
        return jsonify({'transformations': transformations})
    except Exception as e:
        error_msg = f"Error retrieving transformations: {str(e)}"
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': error_msg}), 500

# Auth0 routes
@app.route('/api/auth/login')
def login():
    # Check if already logged in
    if 'logged_in' in session and session['logged_in'] == True:
        add_system_log("Already authenticated user attempting to log in, redirecting to home", "INFO")
        return redirect('/')
    
    # Log the login attempt
    add_system_log("Login attempt initiated, redirecting to Auth0", "INFO")
    
    # Make sure to clear any leftover session data
    session.clear()
    
    # Redirect to Auth0 for authentication
    return auth0.authorize_redirect(
        redirect_uri=url_for('callback', _external=True)
    )

@app.route('/api/auth/callback')
def callback():
    try:
        # Get the auth0 token
        auth0.authorize_access_token()
        resp = auth0.get('userinfo')
        userinfo = resp.json()
        
        add_system_log(f"[AUTH CALLBACK] Auth0 callback received for user: {userinfo.get('email', 'Unknown')}")
        
        # Clear and regenerate session for security
        session.clear()
        
        # Set session to permanent
        session.permanent = True
        
        # Store user info in session
        session['jwt_payload'] = userinfo
        session['profile'] = {
            'user_id': userinfo['sub'],
            'name': userinfo.get('name', ''),
            'picture': userinfo.get('picture', ''),
            'email': userinfo.get('email', '')
        }
        session['logged_in'] = True
        session['auth_time'] = datetime.datetime.now().timestamp()
        
        # Add a session ID for tracking
        session['session_id'] = str(uuid.uuid4())
        
        # Ensure session is saved immediately
        session.modified = True
        
        # Log the successful authentication
        add_system_log(f"[AUTH CALLBACK] User authenticated: {userinfo.get('name', 'Unknown')} ({userinfo.get('email', 'No email')})")
        add_system_log(f"[AUTH CALLBACK] Session data after authentication: {dict(session)}")
        add_system_log(f"[AUTH CALLBACK] Session ID: {id(session)}")
        
        # Handle user record in database if configured
        if users_collection is not None:
            try:
                # Get the user ID from Auth0
                auth0_id = userinfo['sub']
                
                # First try to find user by user_id field (primary identifier)
                user = users_collection.find_one({'user_id': auth0_id})
                
                # If not found, try auth0Id as fallback
                if not user:
                    user = users_collection.find_one({'auth0Id': auth0_id})
                
                timestamp_now = datetime.datetime.now()
                
                # If not, create user record - use field names matching existing data in the db
                if not user:
                    # Generate a username based on email or name to avoid null username issues
                    email = userinfo.get('email', '')
                    name = userinfo.get('name', '')
                    
                    if email:
                        # Try to use email as the basis for username
                        username_base = email.split('@')[0].replace('.', '').lower()
                    elif name:
                        # If no email, use name
                        username_base = name.replace(' ', '').lower()
                    else:
                        # Last resort, generate a random username
                        username_base = f"user_{uuid.uuid4().hex[:8]}"
                    
                    # Add a random number to ensure uniqueness
                    username = f"{username_base}_{random.randint(1000, 9999)}"
                    
                    new_user = {
                        'user_id': auth0_id,             # Primary identifier
                        'auth0Id': auth0_id,             # For backward compatibility
                        'email': userinfo.get('email', ''),
                        'name': userinfo.get('name', ''),
                        'username': username,            # Always set a username
                        'createdAt': timestamp_now,      # Match existing field format in db
                        'lastLogin': timestamp_now,      # Match existing field format in db
                        'lastActive': timestamp_now,     # For admin panel compatibility
                        'subscription': 'FREE',          # Default subscription
                        'usageCount': 0,                 # Initialize usage counter
                        'subscriptionUpdatedAt': timestamp_now,
                        'preferences': {                 # Add user preferences
                            'defaultTone': 'casual',
                            'saveHistory': True
                        }
                    }
                    try:
                        result = users_collection.insert_one(new_user)
                        add_system_log(f"New user created with ID: {result.inserted_id} - {userinfo.get('email', 'No email')}")
                    except Exception as insert_exc:
                        add_system_log(f"CRITICAL: Failed to insert new user: {str(insert_exc)}", "ERROR")
                        return redirect('/login?error=user_creation_failed')
                else:
                    # Update existing user - determine which field names to use based on existing data
                    updates = {
                        'lastLogin': timestamp_now,   # Use matching db format
                        'lastActive': timestamp_now,  # Add for admin panel compatibility
                        # Always ensure user_id exists
                        'user_id': auth0_id
                    }
                    
                    # Update name and email if provided
                    if userinfo.get('name'):
                        updates['name'] = userinfo.get('name')
                    if userinfo.get('email'):
                        updates['email'] = userinfo.get('email')
                        
                    # Add missing fields if they don't exist
                    if 'subscription' not in user:
                        updates['subscription'] = 'FREE'
                    if 'usageCount' not in user:
                        updates['usageCount'] = 0
                    if 'preferences' not in user:
                        updates['preferences'] = {
                            'defaultTone': 'casual',
                            'saveHistory': True
                        }
                    
                    # Check if username is missing and add one if needed
                    if 'username' not in user or not user.get('username'):
                        email = userinfo.get('email', '')
                        name = userinfo.get('name', '')
                        
                        if email:
                            username_base = email.split('@')[0].replace('.', '').lower()
                        elif name:
                            username_base = name.replace(' ', '').lower()
                        
                        # Add a random number to ensure uniqueness
                        username = f"{username_base}_{random.randint(1000, 9999)}"
                        updates['username'] = username
                        add_system_log(f"Added missing username '{username}' to existing user")
                    
                    # Always update with the _id field for consistency
                    users_collection.update_one(
                        {'_id': user['_id']},
                        {'$set': updates}
                    )
                    
                    add_system_log(f"Updated user record for: {userinfo.get('email', 'No email')}")
            except Exception as e:
                add_system_log(f"Error updating user record: {str(e)}", "ERROR")
        
        # Debug log to trace the issue
        add_system_log(f"[AUTH CALLBACK] Authentication complete, redirecting to home page", "INFO")
        
        # Instead of complex JavaScript, just directly redirect to avoid issues
        return redirect("/")
    except Exception as e:
        add_system_log(f"[AUTH CALLBACK] Auth0 callback error: {str(e)}", "ERROR")
        return redirect('/login?error=callback_failed')

@app.route('/api/auth/logout')
def logout():
    # Clear session
    session.clear()
    
    # Log the logout
    add_system_log("User logged out")
    
    # Check if full Auth0 logout is required
    full_logout = request.args.get('full_logout', 'false').lower() == 'true'
    
    if full_logout:
        # Build the logout URL for Auth0
        params = {
            'returnTo': url_for('logout_page', _external=True),
            'client_id': os.getenv('AUTH0_CLIENT_ID')
        }
        logout_url = f"https://{os.getenv('AUTH0_DOMAIN')}/v2/logout?" + urlencode(params)
        return redirect(logout_url)
    else:
        # Just redirect to logout page
        return redirect('/logout')

@app.route('/api/document/generate', methods=['POST'])
def generate_document():
    try:
        # Get request data from JSON
        data = request.get_json() if request.is_json else {}
        text = data.get('text', '')
        format_type = data.get('fileType', 'pdf')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Generate document based on file type
        if format_type == 'pdf':
            return generate_pdf(text)
        elif format_type == 'doc':
            return generate_docx(text)
        elif format_type == 'odt':
            return generate_odt(text)
        else:
            return jsonify({'error': f'Unsupported file type: {format_type}'}), 400
    
    except Exception as e:
        error_msg = f"Error generating document: {str(e)}"
        print(error_msg)
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': 'Failed to generate document'}), 500

def generate_pdf(text):
    """Generate a PDF document from text"""
    try:
        # Create a bytes buffer for the PDF
        buffer = io.BytesIO()
        
        # Create the PDF with ReportLab
        pdf = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Set up font and margins
        pdf.setFont("Helvetica", 12)
        margin = 72  # 1 inch margins
        text_width = width - 2 * margin
        y_position = height - margin
        
        # Process the text and add to PDF
        lines = text.split('\n')
        for line in lines:
            # Use ReportLab's wrap functionality to handle line breaks
            text_object = pdf.beginText(margin, y_position)
            text_object.setFont("Helvetica", 12)
            
            # If it's a blank line, move down
            if not line.strip():
                y_position -= 20
                continue
                
            # Wrap text to fit within margins
            wrapped_text = [line[i:i+80] for i in range(0, len(line), 80)]
            
            for wrap in wrapped_text:
                text_object.textLine(wrap)
                y_position -= 15
                
                # Check if we need a new page
                if y_position < margin:
                    pdf.drawText(text_object)
                    pdf.showPage()
                    pdf.setFont("Helvetica", 12)
                    y_position = height - margin
                    text_object = pdf.beginText(margin, y_position)
                    text_object.setFont("Helvetica", 12)
            
            pdf.drawText(text_object)
        
        # Save the PDF
        pdf.save()
        
        # Move the buffer position to the beginning
        buffer.seek(0)
        
        # Return the PDF as a response
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='mr-wlah-transformed.pdf'
        )
    
    except Exception as e:
        error_msg = f"PDF generation error: {str(e)}"
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': error_msg}), 500

def generate_docx(text):
    """Generate a DOCX document from text"""
    try:
        # Create a new Document
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Mr. Wlah Transformed Text', 0)
        title.alignment = 1  # Center alignment
        
        # Add paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(12)
        
        # Save to a BytesIO object
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return the document as a response
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='mr-wlah-transformed.docx'
        )
    
    except Exception as e:
        error_msg = f"DOCX generation error: {str(e)}"
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': error_msg}), 500

def generate_odt(text):
    """Generate an ODT document from text"""
    try:
        # For ODT format, we'll convert from DOCX
        # First create a DOCX document
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Mr. Wlah Transformed Text', 0)
        title.alignment = 1  # Center alignment
        
        # Add paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(12)
        
        # Save to a temporary file
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_docx.name)
        temp_docx.close()
        
        # Use a third-party conversion like LibreOffice (in a production environment)
        # For this implementation, we'll return a DOCX file with a message
        with open(temp_docx.name, 'rb') as file:
            docx_data = file.read()
        
        # Clean up the temporary file
        os.unlink(temp_docx.name)
        
        # Prepare the response
        buffer = io.BytesIO(docx_data)
        
        # Return the document
        return send_file(
            buffer,
            mimetype='application/vnd.oasis.opendocument.text',
            as_attachment=True,
            download_name='mr-wlah-transformed.odt'
        )
    
    except Exception as e:
        error_msg = f"ODT generation error: {str(e)}"
        add_system_log(error_msg, "ERROR")
        return jsonify({'error': error_msg}), 500

@app.route('/api/auth/status')
def auth_status():
    """Endpoint to check if user is authenticated"""
    is_authenticated = 'logged_in' in session and session['logged_in']
    
    # Add diagnostic logging
    add_system_log(f"[AUTH STATUS] Auth status check - is_authenticated: {is_authenticated}, session ID: {id(session)}")
    if 'profile' in session:
        profile = session.get('profile', {})
        add_system_log(f"[AUTH STATUS] User in session: {profile.get('name', 'Unknown')}")
    
    if is_authenticated:
        profile = session.get('profile', {})
        add_system_log(f"[AUTH STATUS] Returning authenticated status for user: {profile.get('name', 'Unknown')}")
        
        # Ensure session data persists
        session.modified = True
        
        return jsonify({
            'isAuthenticated': True,
            'user': {
                'name': profile.get('name', ''),
                'email': profile.get('email', ''),
                'picture': profile.get('picture', ''),
                'userId': profile.get('user_id', '')  # Include user ID for API calls
            }
        })
    else:
        add_system_log("[AUTH STATUS] Returning unauthenticated status")
        return jsonify({
            'isAuthenticated': False
        })

@app.route('/api/config')
def client_config():
    """Endpoint to provide client-side configuration"""
    return jsonify({
        'auth0': {
            'domain': os.getenv('AUTH0_DOMAIN'),
            'clientId': os.getenv('AUTH0_CLIENT_ID'),
            'audience': os.getenv('AUTH0_AUDIENCE', 'https://api.mrwlah.com'),
            'callbackUrl': f"{request.host_url.rstrip('/')}/api/auth/callback"
        },
        'app': {
            'environment': os.getenv('FLASK_ENV', 'production'),
            'apiBaseUrl': '/api'
        }
    })

@app.route('/api/test/users-collection', methods=['GET'])
def test_users_collection():
    """Test if the 'users' collection is available and writable."""
    try:
        if users_collection is None:
            return jsonify({'success': False, 'error': 'users_collection is None'}), 500
        # Create a unique test user
        test_email = f"test_user_{uuid.uuid4()}@example.com"
        test_user = {
            'auth0Id': f'test-auth0-{uuid.uuid4()}',
            'user_id': f'test-auth0-{uuid.uuid4()}',
            'email': test_email,
            'name': 'Test User',
            'createdAt': datetime.datetime.utcnow(),
            'lastLogin': datetime.datetime.utcnow(),
            'lastActive': datetime.datetime.utcnow(),
            'subscription': 'FREE',
            'usageCount': 0,
            'preferences': {
                'defaultTone': 'casual',
                'saveHistory': True
            }
        }
        # Insert the test user
        result = users_collection.insert_one(test_user)
        inserted_id = result.inserted_id
        # Now delete the test user
        users_collection.delete_one({'_id': inserted_id})
        return jsonify({'success': True, 'message': 'users collection is available and writable'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# Add this function before the main block
def migrate_existing_users():
    """
    Migrate existing users to ensure they have the required fields.
    This function should run once at application startup.
    """
    if users_collection is None:
        add_system_log("No database connection, skipping user migration", "WARNING")
        return False
    
    try:
        # Find all users
        users = list(users_collection.find())
        
        if not users:
            add_system_log("No users to migrate", "INFO")
            return True
        
        add_system_log(f"Found {len(users)} users to check for migration", "INFO")
        
        migrated_count = 0
        for user in users:
            updates = {}
            
            # Ensure user_id field exists and is prioritized
            # First check if user has an auth0Id but no user_id
            if 'auth0Id' in user and 'user_id' not in user:
                updates['user_id'] = user['auth0Id']
            # Also handle auth0_id field (with underscore)
            elif 'auth0_id' in user and 'user_id' not in user:
                updates['user_id'] = user['auth0_id']
            # If neither field exists, create a placeholder user_id
            elif 'user_id' not in user:
                updates['user_id'] = str(user['_id'])  # Use MongoDB _id as a fallback
                updates['auth0Id'] = str(user['_id'])  # Set auth0Id too for consistency
            
            # Ensure other required fields
            if 'lastActive' not in user and 'lastLogin' in user:
                updates['lastActive'] = user['lastLogin']
            
            if 'subscription' not in user:
                updates['subscription'] = 'FREE'
            
            if 'usageCount' not in user:
                updates['usageCount'] = 0
            
            if 'preferences' not in user:
                updates['preferences'] = {
                    'defaultTone': 'casual',
                    'saveHistory': True
                }
            
            # If we need to update the user
            if updates:
                users_collection.update_one(
                    {'_id': user['_id']},
                    {'$set': updates}
                )
                migrated_count += 1
        
        if migrated_count > 0:
            add_system_log(f"Migrated {migrated_count} users to new format", "INFO")
        else:
            add_system_log("All users are already in the correct format", "INFO")
        
        return True
    except Exception as e:
        add_system_log(f"Error migrating users: {str(e)}", "ERROR")
        return False

if __name__ == '__main__':
    try:
        app.jinja_env.auto_reload = True
        app.config['TEMPLATES_AUTO_RELOAD'] = True
        
        # Get port and environment from .env
        port = int(os.getenv('PORT', 3000))
        debug_mode = os.getenv('NODE_ENV') == 'development'
        
        # Log startup information
        start_msg = f"Starting Mr. Wlah application on port {port}"
        if debug_mode:
            start_msg += " in DEBUG mode"
        add_system_log(start_msg, "INFO")
        
        # Log database status
        if users_collection is not None:
            db_status = f"Connected to {mongo_db} database"
            if transformations_collection is not None:
                try:
                    transform_count = transformations_collection.count_documents({})
                    db_status += f" with {transform_count} transformations"
                except Exception as e:
                    db_status += f" (error counting transformations: {str(e)})"
            add_system_log(db_status, "INFO")
            
            # Migrate existing users to new format
            migrate_existing_users()
        else:
            add_system_log("Running without database connection", "WARNING")
        
        # Start the Flask app
        app.run(host='0.0.0.0', port=port, debug=debug_mode)
    except Exception as e:
        error_msg = f"Application startup error: {str(e)}"
        print(f"❌ {error_msg}")
        add_system_log(error_msg, "ERROR")
        sys.exit(1) 